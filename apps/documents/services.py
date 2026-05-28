import json
import os
import tempfile

from django.core.files import File
from django.db import transaction
from docx import Document as DocxDocument

from apps.documents.models import Documents
from apps.logs.services import write_operation_log
from apps.processes.models import ProcessSteps
from apps.requisites.models import RequisiteLinks, Requisites, RequisiteValues
from apps.requisites.services import validate_document_requisites
from apps.templates_cfg.docx_table import (
    apply_table_rows_to_docx,
    expand_duplicate_table_rows_in_html,
    find_table_template_anchor,
    parse_table_anchor_json,
    replace_placeholder_values_in_text,
)


class DocumentAutofillService:
    def build_document_form(self, document_type_id, process_id, object_id, user):
        requisites = Requisites.objects.filter(document_type_id_id=document_type_id).order_by("requisite_id")
        links = RequisiteLinks.objects.select_related("source_requisite_id", "target_requisite_id").filter(
            target_requisite_id__in=requisites
        )
        link_map = {link.target_requisite_id_id: link for link in links}
        fields = []
        warnings = []

        for req in requisites:
            field_payload = {
                "requisite_id": req.requisite_id,
                "name": req.name,
                "data_type": req.data_type,
                "is_required": req.is_required,
                "field_kind": req.field_kind,
                "placeholder_key": req.placeholder_key,
                "table_column": getattr(req, "is_table_column", False),
                "value": "",
                "source": "manual",
                "readonly": False,
                "source_document": "",
            }
            link = link_map.get(req.requisite_id)
            if link:
                source_req = link.source_requisite_id
                source_doc_qs = Documents.objects.filter(document_type_id=source_req.document_type_id)
                if link.inheritance_rule in {"copy_latest_value", "copy_from_same_process"}:
                    source_doc_qs = source_doc_qs.filter(process_id_id=process_id)
                if link.inheritance_rule in {"copy_latest_value", "copy_from_same_object"}:
                    source_doc_qs = source_doc_qs.filter(object_id_id=object_id)
                if link.inheritance_rule == "copy_from_same_process_and_object":
                    source_doc_qs = source_doc_qs.filter(process_id_id=process_id, object_id_id=object_id)
                source_doc_qs = source_doc_qs.order_by("-created_at")
                source_doc = source_doc_qs.first()
                source_value = None
                if source_doc:
                    source_value = RequisiteValues.objects.filter(
                        document_id=source_doc,
                        requisite_id=source_req,
                    ).first()
                if source_value and source_value.value:
                    field_payload["value"] = source_value.value
                    field_payload["source"] = "auto"
                    field_payload["readonly"] = True
                    field_payload["source_document"] = f"{source_doc.document_type_id.name} #{source_doc.document_id}"
                    write_operation_log(
                        user=user,
                        document=source_doc,
                        operation_result=f"AUTOFILL_FIELD:{source_req.name}->{req.name}",
                    )
                else:
                    warnings.append(
                        f"Поле '{req.name}' не было найдено в связанных документах. "
                        "Заполните вручную или проверьте настройки связей."
                    )
            fields.append(field_payload)

        return {"fields": fields, "warnings": warnings}


def build_generated_html(document):
    dt = document.document_type_id
    template_html = dt.template_html or ""
    if getattr(dt, "has_table_template", False):
        try:
            rows = json.loads(document.table_rows_json or "[]")
        except json.JSONDecodeError:
            rows = []
        rendered_html = expand_duplicate_table_rows_in_html(template_html, rows if isinstance(rows, list) else [])
    else:
        rendered_html = template_html
    value_by_req = {
        rv.requisite_id_id: rv.value
        for rv in RequisiteValues.objects.select_related("requisite_id").filter(document_id=document)
    }
    flat: dict[str, str] = {}
    for req in Requisites.objects.filter(document_type_id=dt):
        if getattr(req, "is_table_column", False):
            continue
        if req.placeholder_key:
            flat[req.placeholder_key] = str(value_by_req.get(req.requisite_id, ""))
    rendered_html = replace_placeholder_values_in_text(rendered_html, flat)
    return rendered_html


def _replace_keys_in_paragraph(paragraph, values_by_key: dict) -> None:
    text = paragraph.text
    new = replace_placeholder_values_in_text(text, values_by_key)
    if new != text:
        paragraph.text = new


def build_generated_docx(document):
    template_file = document.document_type_id.template_file
    if not template_file:
        return None
    source_path = template_file.path
    if not source_path.lower().endswith(".docx"):
        return None
    doc = DocxDocument(source_path)
    dt = document.document_type_id
    try:
        rows_data = json.loads(document.table_rows_json or "[]")
    except json.JSONDecodeError:
        rows_data = []
    if not isinstance(rows_data, list):
        rows_data = []
    if getattr(dt, "has_table_template", False) and rows_data:
        anchor = parse_table_anchor_json(getattr(dt, "table_anchor_json", None))
        if not anchor:
            anchor = find_table_template_anchor(doc)
        if anchor:
            apply_table_rows_to_docx(doc, anchor, rows_data)
    values_by_key = {
        req.placeholder_key: rv.value
        for rv in RequisiteValues.objects.select_related("requisite_id")
        .filter(document_id=document)
        for req in [rv.requisite_id]
        if req.placeholder_key and not getattr(req, "is_table_column", False)
    }
    def _replace_in_table_recursive(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_keys_in_paragraph(paragraph, values_by_key)
                for inner in cell.tables:
                    _replace_in_table_recursive(inner)

    for paragraph in doc.paragraphs:
        _replace_keys_in_paragraph(paragraph, values_by_key)
    for table in doc.tables:
        _replace_in_table_recursive(table)
    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc.save(tmp_path)
        with open(tmp_path, "rb") as generated:
            filename = f"document_{document.document_id}.docx"
            document.generated_file.save(filename, File(generated), save=False)
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass
    return document.generated_file


@transaction.atomic
def generate_document(document, values_payload, user):
    for item in values_payload:
        req = Requisites.objects.get(requisite_id=item["requisite_id"])
        RequisiteValues.objects.update_or_create(
            document_id=document,
            requisite_id=req,
            defaults={"value": str(item.get("value", ""))},
        )
    errors = validate_document_requisites(document)
    if errors:
        raise ValueError("; ".join(errors))
    document.status = "generated"
    document.generated_html = build_generated_html(document)
    build_generated_docx(document)
    document.save(update_fields=["status", "generated_html", "generated_file"])
    first_step = ProcessSteps.objects.filter(process_id=document.process_id).order_by("step_order").first()
    write_operation_log(user=user, document=document, step=first_step, operation_result="DOCUMENT_GENERATED")
    return document


@transaction.atomic
def move_document(document, user, result):
    steps = ProcessSteps.objects.filter(process_id=document.process_id).order_by("step_order")
    last_log = document.operation_logs.select_related("step_id").first()
    if not last_log:
        step = steps.first()
    else:
        step = steps.filter(step_order__gt=last_log.step_id.step_order).first()
    if not step:
        document.status = "archived"
        document.save(update_fields=["status"])
        return {"status": document.status, "step": None}
    document.status = "generated"
    document.save(update_fields=["status"])
    write_operation_log(user=user, document=document, step=step, operation_result=result)
    return {"status": document.status, "step": step.step_id}
