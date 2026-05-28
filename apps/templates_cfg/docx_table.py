"""
DOCX: извлечение плейсхолдеров из абзацев и ячеек таблиц; поиск пары строк-шаблона таблицы; заполнение и размножение строк.

Важно: плейсхолдеры во вложенных таблицах (таблица внутри ячейки) не попадают в ``cell.text`` верхнего уровня —
используется рекурсивный обход ``cell.tables`` и единый список таблиц в глубину для якоря и подстановки.
"""
from __future__ import annotations

import html as html_module
import json
import re
import zipfile
from copy import deepcopy
from typing import Any

from docx import Document as DocxDocument
from docx.table import _Row

_PLACEHOLDER_RE = re.compile(
    r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}|\{\(\s*([a-zA-Z0-9_]+)\s*\)\}|\{\?\s*([a-zA-Z0-9_]+)\s*\}"
)


def replace_placeholder_values_in_text(text: str, mapping: dict[str, str]) -> str:
    """
    Подставляет значения вместо плейсхолдеров в строке (DOCX/HTML).
    Допускает пробелы вокруг ключа, как в Word: ``{{ pod }}``, ``{{pod}}``.
    Значения подставляются как есть (без интерпретации ``\\`` как ссылок на группы).
    """
    if not text or not mapping:
        return text
    out = text
    for key, val in mapping.items():
        if not key:
            continue
        sval = str(val if val is not None else "")
        ek = re.escape(key)
        out = re.sub(rf"\{{\{{\s*{ek}\s*\}}\}}", lambda m, sv=sval: sv, out)
        out = re.sub(rf"\{{\(\s*{ek}\s*\)\}}", lambda m, sv=sval: sv, out)
        out = re.sub(rf"\{{\?\s*{ek}\s*\}}", lambda m, sv=sval: sv, out)
    return out


def _keys_from_text(text: str) -> frozenset[str]:
    keys: set[str] = set()
    for m in _PLACEHOLDER_RE.finditer(text or ""):
        for g in m.groups():
            if g:
                keys.add(g)
    return frozenset(keys)


def _cell_text_recursive(cell) -> str:
    """Текст ячейки: абзацы + рекурсивно все вложенные таблицы (Word часто кладёт сетку внутрь ячейки)."""
    bits: list[str] = []
    for p in cell.paragraphs:
        t = p.text or ""
        if t.strip():
            bits.append(t)
    for tbl in cell.tables:
        bits.append(_table_text_recursive(tbl))
    return "\n".join(bits)


def _table_text_recursive(table) -> str:
    lines: list[str] = []
    for row in table.rows:
        lines.append(_row_text(row))
    return "\n".join(lines)


def _row_text(row) -> str:
    parts = []
    for cell in row.cells:
        parts.append(_cell_text_recursive(cell))
    return "\n".join(parts)


def _enumerate_tables_depth_first(doc) -> list:
    """
    Все объекты Table в теле документа в порядке DFS (как в Word: сначала таблица, затем вложенные в её ячейках).
    ``table_idx`` в якоре — индекс в этом списке (не только ``doc.tables`` верхнего уровня).
    """
    out: list = []

    def visit_table(t) -> None:
        out.append(t)
        for row in t.rows:
            for cell in row.cells:
                for nt in cell.tables:
                    visit_table(nt)

    for t in doc.tables:
        visit_table(t)
    return out


def _opc_flat_text_for_placeholder_scan(file_path: str) -> str:
    """
    Плоский текст из OPC XML (document/header/footer): подстраховка, если фрагменты размечены нестандартно
    или плейсхолдер разбит так, что высокоуровневый API даёт неполную картину.
    """
    chunks: list[str] = []
    try:
        with zipfile.ZipFile(file_path, "r") as zf:
            names = sorted(zf.namelist())
            for name in names:
                if not name.startswith("word/") or not name.endswith(".xml"):
                    continue
                leaf = name.rsplit("/", 1)[-1]
                include = leaf in ("document.xml", "footnotes.xml", "endnotes.xml")
                include = include or (leaf.startswith("header") and leaf.endswith(".xml"))
                include = include or (leaf.startswith("footer") and leaf.endswith(".xml"))
                if not include:
                    continue
                try:
                    data = zf.read(name).decode("utf-8", errors="replace")
                except (OSError, KeyError, UnicodeError):
                    continue
                data = re.sub(r"<w:instrText[^>]*>.*?</w:instrText>", " ", data, flags=re.DOTALL)
                plain = re.sub(r"<[^>]+>", " ", data)
                plain = html_module.unescape(plain)
                plain = re.sub(r"\s+", " ", plain)
                chunks.append(plain)
    except (zipfile.BadZipFile, OSError):
        pass
    return "\n".join(chunks)


def _keys_in_row(row) -> frozenset[str]:
    return _keys_from_text(_row_text(row))


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    text = paragraph.text
    new = replace_placeholder_values_in_text(text, mapping)
    if new != text:
        paragraph.text = new


def _fill_row_cells(row, mapping: dict[str, str]) -> None:
    for cell in row.cells:
        for p in cell.paragraphs:
            _replace_in_paragraph(p, mapping)


def full_docx_text_for_placeholders(file_path: str) -> str:
    """Весь текст DOCX: абзацы тела + рекурсивно все таблицы (включая вложенные)."""
    doc = DocxDocument(file_path)
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        parts.append(_table_text_recursive(table))
    opc = _opc_flat_text_for_placeholder_scan(file_path)
    if opc.strip():
        parts.append(opc)
    return "\n".join(parts)


def extract_placeholder_keys_from_docx(file_path: str) -> list[str]:
    """Уникальные ключи плейсхолдеров по всему документу."""
    text = full_docx_text_for_placeholders(file_path)
    keys: set[str] = set()
    for m in _PLACEHOLDER_RE.finditer(text):
        for g in m.groups():
            if g:
                keys.add(g)
    return sorted(keys)


def find_table_template_anchor(doc) -> dict[str, int] | None:
    """
    Ищет первую пару подряд идущих строк таблицы с одинаковым непустым набором плейсхолдеров
    (сигнал «шаблон строки таблицы»).
    Возвращает {"table_idx", "row0", "row1"} или None.
    """
    tables = _enumerate_tables_depth_first(doc)
    for ti, table in enumerate(tables):
        rows = table.rows
        for i in range(len(rows) - 1):
            k0 = _keys_in_row(rows[i])
            k1 = _keys_in_row(rows[i + 1])
            if k0 and k0 == k1:
                return {"table_idx": ti, "row0": i, "row1": i + 1}
    return None


def expand_duplicate_table_rows_in_html(html: str, rows_data: list[dict[str, Any]]) -> str:
    """
    В HTML-предпросмотре заменяет первую пару подряд идущих <tr> с одинаковым набором плейсхолдеров
    на N строк с подстановкой значений из rows_data (шаблон строки — первый <tr> пары).
    """
    if not rows_data or not html:
        return html
    tr_re = re.compile(r"<tr[^>]*>.*?</tr>", re.IGNORECASE | re.DOTALL)
    trs = tr_re.findall(html)
    if len(trs) < 2:
        return html

    def tr_plain(tr_html: str) -> str:
        return re.sub(r"<[^>]+>", " ", tr_html)

    for i in range(len(trs) - 1):
        k0 = _keys_from_text(tr_plain(trs[i]))
        k1 = _keys_from_text(tr_plain(trs[i + 1]))
        if not k0 or k0 != k1:
            continue
        template_tr = trs[i]

        def subst(tr_tpl: str, row: dict[str, Any]) -> str:
            row_str = {str(k): str(v if v is not None else "") for k, v in row.items()}
            return replace_placeholder_values_in_text(tr_tpl, row_str)

        new_block = "".join(subst(template_tr, row) for row in rows_data)
        old_block = trs[i] + trs[i + 1]
        return html.replace(old_block, new_block, 1)
    return html


def parse_table_anchor_json(raw: str | None) -> dict[str, int] | None:
    if not raw or not str(raw).strip():
        return None
    try:
        d = json.loads(raw)
    except json.JSONDecodeError:
        return None
    if not isinstance(d, dict):
        return None
    try:
        return {
            "table_idx": int(d["table_idx"]),
            "row0": int(d["row0"]),
            "row1": int(d["row1"]),
        }
    except (KeyError, TypeError, ValueError):
        return None


def apply_table_rows_to_docx(doc, anchor: dict[str, int], rows_data: list[dict[str, Any]]) -> None:
    """
    Заполняет две строки-шаблона данными rows_data[0], rows_data[1];
    при len(rows_data) > 2 дублирует структуру второй строки и заполняет последующие строки.
    """
    if not rows_data:
        return
    ti = anchor["table_idx"]
    r0 = anchor["row0"]
    r1 = anchor["row1"]
    tables = _enumerate_tables_depth_first(doc)
    if ti < 0 or ti >= len(tables):
        return
    table = tables[ti]
    template_tr = deepcopy(table.rows[r1]._tr)
    row0 = table.rows[r0]
    row1 = table.rows[r1]
    key_union = sorted(_keys_in_row(row0) | _keys_in_row(row1))
    blank = {k: "" for k in key_union}
    _fill_row_cells(row0, rows_data[0])
    if len(rows_data) > 1:
        _fill_row_cells(row1, rows_data[1])
    else:
        _fill_row_cells(row1, blank)
    anchor_tr = row1._tr
    cur_tr = anchor_tr
    for idx in range(2, len(rows_data)):
        new_tr = deepcopy(template_tr)
        cur_tr.addnext(new_tr)
        cur_tr = new_tr
        new_row = _Row(new_tr, table)
        _fill_row_cells(new_row, rows_data[idx])
