import os
import tempfile

from django.test import SimpleTestCase
from docx import Document as DocxDocument

from apps.templates_cfg.docx_table import (
    extract_placeholder_keys_from_docx,
    find_table_template_anchor,
    full_docx_text_for_placeholders,
    replace_placeholder_values_in_text,
)


class DocxTablePlaceholderTests(SimpleTestCase):
    def test_replace_placeholder_allows_spaces(self):
        self.assertEqual(
            replace_placeholder_values_in_text("A {{ pod }} B", {"pod": "X"}),
            "A X B",
        )
        self.assertEqual(
            replace_placeholder_values_in_text("{{name2}} and {{ name2 }}", {"name2": "Z"}),
            "Z and Z",
        )

    def test_nested_table_placeholders_visible(self):
        doc = DocxDocument()
        outer = doc.add_table(rows=1, cols=1)
        outer_cell = outer.rows[0].cells[0]
        inner = outer_cell.add_table(rows=2, cols=2)
        inner.rows[0].cells[0].text = "{{a}}"
        inner.rows[0].cells[1].text = "{{b}}"
        inner.rows[1].cells[0].text = "{{a}}"
        inner.rows[1].cells[1].text = "{{b}}"
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)
            keys = extract_placeholder_keys_from_docx(path)
            self.assertEqual(sorted(keys), ["a", "b"])
            blob = full_docx_text_for_placeholders(path)
            self.assertIn("{{a}}", blob)
            self.assertIn("{{b}}", blob)
            d2 = DocxDocument(path)
            anchor = find_table_template_anchor(d2)
            self.assertIsNotNone(anchor)
            self.assertEqual(anchor["row0"], 0)
            self.assertEqual(anchor["row1"], 1)
            # Вложенная таблица — второй объект в DFS после внешней
            self.assertEqual(anchor["table_idx"], 1)
        finally:
            os.unlink(path)

    def test_flat_table_two_rows_anchor(self):
        doc = DocxDocument()
        t = doc.add_table(rows=3, cols=2)
        t.rows[1].cells[0].text = "{{x}}"
        t.rows[1].cells[1].text = "{{y}}"
        t.rows[2].cells[0].text = "{{x}}"
        t.rows[2].cells[1].text = "{{y}}"
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)
            self.assertEqual(extract_placeholder_keys_from_docx(path), ["x", "y"])
            d2 = DocxDocument(path)
            anchor = find_table_template_anchor(d2)
            self.assertIsNotNone(anchor)
            self.assertEqual(anchor["table_idx"], 0)
            self.assertEqual(anchor["row0"], 1)
            self.assertEqual(anchor["row1"], 2)
        finally:
            os.unlink(path)
