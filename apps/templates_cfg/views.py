import json
import re

from django.http import Http404
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.views import View
from docx import Document as DocxDocument
from rest_framework.viewsets import ModelViewSet

from apps.documents.models import Documents
from apps.logs.services import write_operation_log
from apps.users.permissions import RoleBasedPermission
from apps.users.ui_permissions import LoginRequiredRoleMixin
from apps.requisites.models import Requisites
from apps.bpm.models import FieldSourceRule, ProcessDocumentTemplate
from .docx_table import (
    extract_placeholder_keys_from_docx,
    find_table_template_anchor,
    replace_placeholder_values_in_text,
)
from .models import DocumentTypes, ProductionObjects
from .serializers import DocumentTypeSerializer, ProductionObjectSerializer


class SpecialistPermission(RoleBasedPermission):
    allowed_roles = {"ADMIN", "SPECIALIST"}


class DocumentTypeViewSet(ModelViewSet):
    queryset = DocumentTypes.objects.all()
    serializer_class = DocumentTypeSerializer
    permission_classes = [SpecialistPermission]

    def perform_create(self, serializer):
        doc_type = serializer.save()
        write_operation_log(user=self.request.enterprise_user, operation_result=f"DOCUMENT_TYPE_CREATED:{doc_type.name}")


class ProductionObjectViewSet(ModelViewSet):
    queryset = ProductionObjects.objects.all()
    serializer_class = ProductionObjectSerializer
    permission_classes = [SpecialistPermission]


class UIProductionObjectsView(LoginRequiredRoleMixin, View):
    allowed_roles = {"ADMIN", "SPECIALIST"}

    def get(self, request):
        edit_id = request.GET.get("edit_id")
        edit_object = ProductionObjects.objects.filter(object_id=edit_id).first() if edit_id else None
        objects = ProductionObjects.objects.all().order_by("object_type", "name")
        return render(request, "ui/specialist/production_objects.html", {"objects": objects, "edit_object": edit_object})

    def post(self, request):
        action = request.POST.get("action", "create")
        if action == "delete":
            obj = ProductionObjects.objects.get(object_id=request.POST["object_id"])
            write_operation_log(user=request.enterprise_user, operation_result=f"PRODUCTION_OBJECT_DELETED:{obj.name}")
            obj.delete()
            return redirect("ui-production-objects")
        if action == "update":
            obj = ProductionObjects.objects.get(object_id=request.POST["object_id"])
            obj.object_type = request.POST.get("object_type", "").strip() or obj.object_type
            obj.name = request.POST.get("name", "").strip() or obj.name
            obj.save()
            write_operation_log(user=request.enterprise_user, operation_result=f"PRODUCTION_OBJECT_UPDATED:{obj.name}")
            return redirect("ui-production-objects")
        obj = ProductionObjects.objects.create(
            object_type=request.POST.get("object_type", "").strip() or "Объект",
            name=request.POST.get("name", "").strip() or "Без названия",
        )
        write_operation_log(user=request.enterprise_user, operation_result=f"PRODUCTION_OBJECT_CREATED:{obj.name}")
        return redirect("ui-production-objects")


def extract_placeholders_from_docx(file_path):
    """Уникальные ключи плейсхолдеров по абзацам и ячейкам таблиц DOCX."""
    return extract_placeholder_keys_from_docx(file_path)


def sync_doc_type_table_fields_after_docx_upload(doc_type: DocumentTypes, has_table: bool) -> str | None:
    """
    Обновляет template_html, флаги таблицы, якорь строк; создаёт/обновляет реквизиты по ключам из DOCX.
    Возвращает сообщение об ошибке или None при успехе.
    """
    if not doc_type.template_file:
        return "Нет файла шаблона"
    path = doc_type.template_file.path
    if not str(path).lower().endswith(".docx"):
        return None
    placeholders = extract_placeholder_keys_from_docx(path)
    doc_type.template_html = docx_to_template_html(path)
    doc_type.has_table_template = bool(has_table)
    if has_table:
        d = DocxDocument(path)
        anchor = find_table_template_anchor(d)
        if not anchor:
            return (
                "Не найдена пара подряд идущих строк таблицы с одинаковым непустым набором плейсхолдеров. "
                "См. руководство: для шаблона с таблицей нужны две одинаковые строки ключей."
            )
        doc_type.table_anchor_json = json.dumps(anchor, ensure_ascii=False)
    else:
        doc_type.table_anchor_json = ""
    doc_type.save(update_fields=["template_html", "has_table_template", "table_anchor_json"])
    for ph in placeholders:
        req, _ = Requisites.objects.get_or_create(
            document_type_id=doc_type,
            placeholder_key=ph,
            defaults={
                "name": ph,
                "data_type": "text",
                "is_required": True,
                "field_kind": "variable",
                "is_table_column": bool(has_table),
            },
        )
        if req.is_table_column != bool(has_table):
            req.is_table_column = bool(has_table)
            req.save(update_fields=["is_table_column"])
    return None


def docx_to_template_html(file_path):
    doc = DocxDocument(file_path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(f"<p>{p.text}</p>")

    def _cell_inner_html(cell) -> str:
        inner: list[str] = []
        for cp in cell.paragraphs:
            if cp.text.strip():
                inner.append(f"<p>{cp.text}</p>")
        for inner_tbl in cell.tables:
            inner.append(_table_to_html(inner_tbl))
        return "".join(inner) if inner else "&nbsp;"

    def _table_to_html(table) -> str:
        rows_html: list[str] = ["<table border='1' cellpadding='6' cellspacing='0' style='border-collapse:collapse;width:100%;margin:10px 0;'>"]
        for row in table.rows:
            rows_html.append("<tr>")
            for cell in row.cells:
                rows_html.append(f"<td>{_cell_inner_html(cell)}</td>")
            rows_html.append("</tr>")
        rows_html.append("</table>")
        return "".join(rows_html)

    for table in doc.tables:
        parts.append(_table_to_html(table))
    return "".join(parts)


class UITemplateListView(LoginRequiredRoleMixin, View):
    allowed_roles = {"ADMIN", "SPECIALIST"}

    def get(self, request):
        documents = Documents.objects.select_related("document_type_id", "process_id", "object_id", "user_id").all()[:200]
        return render(request, "ui/specialist/documents_registry.html", {"documents": documents})


class UIDocumentTypeCrudView(LoginRequiredRoleMixin, View):
    allowed_roles = {"ADMIN", "SPECIALIST"}

    def get(self, request):
        types = DocumentTypes.objects.all().order_by("name")
        q = request.GET.get("q", "").strip()
        if q:
            types = types.filter(name__icontains=q)
        edit_id = request.GET.get("edit_id")
        edit_type = DocumentTypes.objects.filter(document_type_id=edit_id).first() if edit_id else None
        return render(
            request,
            "ui/specialist/document_types.html",
            {"types": types, "edit_type": edit_type, "search_q": q},
        )

    def post(self, request):
        action = request.POST.get("action", "")
        if action == "update_meta":
            doc_type = DocumentTypes.objects.get(document_type_id=request.POST["document_type_id"])
            name = request.POST.get("name", "").strip()
            if name:
                doc_type.name = name
            doc_type.description = request.POST.get("description", "").strip()
            has_table = request.POST.get("has_table_template") == "on"
            doc_type.save(update_fields=["name", "description"])
            if doc_type.template_file and str(doc_type.template_file.name).lower().endswith(".docx"):
                err = sync_doc_type_table_fields_after_docx_upload(doc_type, has_table)
                if err:
                    q = request.GET.get("q", "").strip()
                    types = DocumentTypes.objects.all().order_by("name")
                    if q:
                        types = types.filter(name__icontains=q)
                    return render(
                        request,
                        "ui/specialist/document_types.html",
                        {
                            "types": types,
                            "edit_type": doc_type,
                            "search_q": q,
                            "error": err,
                        },
                    )
            else:
                if has_table:
                    q = request.GET.get("q", "").strip()
                    types = DocumentTypes.objects.all().order_by("name")
                    if q:
                        types = types.filter(name__icontains=q)
                    return render(
                        request,
                        "ui/specialist/document_types.html",
                        {
                            "types": types,
                            "edit_type": doc_type,
                            "search_q": q,
                            "error": "Режим таблицы требует загруженного DOCX. Откройте «Поля» и загрузите шаблон, затем включите флажок снова.",
                        },
                    )
                doc_type.has_table_template = False
                doc_type.table_anchor_json = ""
                doc_type.save(update_fields=["has_table_template", "table_anchor_json"])
                Requisites.objects.filter(document_type_id=doc_type).update(is_table_column=False)
            write_operation_log(
                user=request.enterprise_user,
                operation_result=f"DOCUMENT_TYPE_UPDATED:{doc_type.name}",
            )
            return redirect("ui-document-types")
        if action == "delete":
            doc_type = DocumentTypes.objects.get(document_type_id=request.POST["document_type_id"])
            write_operation_log(user=request.enterprise_user, operation_result=f"DOCUMENT_TYPE_DELETED:{doc_type.name}")
            doc_type.delete()
            return redirect("ui-document-types")
        return redirect("ui-document-types")


class UITemplateCreateView(LoginRequiredRoleMixin, View):
    allowed_roles = {"ADMIN", "SPECIALIST"}

    def get(self, request):
        return render(request, "ui/specialist/template_create.html")

    def post(self, request):
        title = request.POST.get("name", "").strip()
        description = request.POST.get("description", "").strip()
        template_file = request.FILES.get("template_file")
        if not title or not template_file:
            return render(request, "ui/specialist/template_create.html", {"error": "Заполните название и загрузите DOCX"})
        if DocumentTypes.objects.filter(name__iexact=title).exists():
            return render(request, "ui/specialist/template_create.html", {"error": "Тип документа с таким названием уже существует"})
        if template_file.name.lower().endswith(".doc") and not template_file.name.lower().endswith(".docx"):
            return render(
                request,
                "ui/specialist/template_create.html",
                {"warning": "Для корректной обработки используйте формат DOCX"},
            )
        has_table = request.POST.get("has_table_template") == "on"
        doc_type = DocumentTypes.objects.create(name=title, description=description, template_file=template_file)
        err = sync_doc_type_table_fields_after_docx_upload(doc_type, has_table)
        if err:
            doc_type.delete()
            return render(request, "ui/specialist/template_create.html", {"error": err})
        placeholders = extract_placeholders_from_docx(doc_type.template_file.path)
        write_operation_log(user=request.enterprise_user, operation_result=f"TEMPLATE_UPLOADED:{doc_type.name}")
        write_operation_log(user=request.enterprise_user, operation_result=f"DOCX_FIELDS_EXTRACTED:{len(placeholders)}")
        return redirect("ui-template-configure", document_type_id=doc_type.document_type_id)


class UITemplateConfigureFieldsView(LoginRequiredRoleMixin, View):
    allowed_roles = {"ADMIN", "SPECIALIST"}

    def get(self, request, document_type_id):
        doc_type = DocumentTypes.objects.get(document_type_id=document_type_id)
        requisites = Requisites.objects.filter(document_type_id=doc_type).order_by("requisite_id")
        return render(
            request,
            "ui/specialist/template_configure_fields.html",
            {
                "doc_type": doc_type,
                "requisites": requisites,
            },
        )

    def post(self, request, document_type_id):
        doc_type = DocumentTypes.objects.get(document_type_id=document_type_id)
        if request.POST.get("action") == "add_manual":
            Requisites.objects.create(
                document_type_id=doc_type,
                placeholder_key=request.POST.get("manual_placeholder", "").strip(),
                name=request.POST.get("manual_name", "").strip() or "Новый реквизит",
                data_type=request.POST.get("manual_data_type", "text"),
                is_required=request.POST.get("manual_required") == "on",
                field_kind="variable",
                is_table_column=False,
            )
            write_operation_log(user=request.enterprise_user, operation_result=f"REQUISITE_CREATED_MANUAL:{doc_type.name}")
            return redirect("ui-template-configure", document_type_id=doc_type.document_type_id)
        requisites = Requisites.objects.filter(document_type_id=doc_type)
        for req in requisites:
            req.name = request.POST.get(f"name_{req.requisite_id}", req.name)
            req.data_type = request.POST.get(f"data_type_{req.requisite_id}", req.data_type)
            req.field_kind = "variable"
            req.is_required = request.POST.get(f"is_required_{req.requisite_id}") == "on"
            req.save()
            write_operation_log(user=request.enterprise_user, operation_result=f"REQUISITE_UPDATED:{req.name}")
        url = reverse("ui-template-preview", kwargs={"document_type_id": doc_type.document_type_id})
        return redirect(f"{url}?simple=1")


class UITemplatePreviewView(LoginRequiredRoleMixin, View):
    allowed_roles = {"ADMIN", "SPECIALIST"}

    def get(self, request, document_type_id):
        doc_type = DocumentTypes.objects.get(document_type_id=document_type_id)
        if not (doc_type.template_html or "").strip() and doc_type.template_file and doc_type.template_file.name.lower().endswith(".docx"):
            doc_type.template_html = docx_to_template_html(doc_type.template_file.path)
            doc_type.save(update_fields=["template_html"])
        requisites = Requisites.objects.filter(document_type_id=doc_type)
        if not (doc_type.template_html or "").strip() and requisites.exists():
            doc_type.template_html = "".join([f"<p>{{{{{r.placeholder_key or r.name}}}}}</p>" for r in requisites])
            doc_type.save(update_fields=["template_html"])

        pdt_raw = request.GET.get("pdt_id")
        bpt_raw = request.GET.get("bpt_id")
        rules_by_req = {}
        bpm_step_preview = False
        if pdt_raw and bpt_raw:
            pdt = get_object_or_404(
                ProcessDocumentTemplate.objects.select_related("document_type"),
                pdt_id=int(pdt_raw),
                business_process_template_id=int(bpt_raw),
            )
            if pdt.document_type_id != doc_type.document_type_id:
                raise Http404("Этап шаблона не соответствует типу документа")
            rules_by_req = {
                r.requisite_id: r
                for r in FieldSourceRule.objects.filter(process_document_template=pdt)
            }
            bpm_step_preview = True

        simple = (request.GET.get("simple") == "1") and not bpm_step_preview
        html = doc_type.template_html or ""

        def replace_req_placeholder(key: str, kind: str, badge: str) -> None:
            nonlocal html
            replacement = (
                f"<span class='ph {kind}' title='{badge}'>&#123;&#123;{key}&#125;&#125;</span>"
            )
            html = replace_placeholder_values_in_text(html, {key: replacement})

        if bpm_step_preview:
            for req in requisites:
                key = req.placeholder_key or req.name
                rule = rules_by_req.get(req.requisite_id)
                if rule and rule.source_type in (
                    FieldSourceRule.SOURCE_DICTIONARY,
                    FieldSourceRule.SOURCE_PREVIOUS_DOCUMENT,
                ):
                    kind, badge = (
                        "constant",
                        "Постоянный реквизит — автоподстановка из справочника или предыдущего документа",
                    )
                else:
                    kind, badge = "variable", "Переменный реквизит — ручной ввод оператором на этом шаге"
                replace_req_placeholder(key, kind, badge)
        elif not simple:
            for req in requisites:
                key = req.placeholder_key or req.name
                replace_req_placeholder(
                    key,
                    "variable",
                    "В каталоге типов поле переменное; автозаполнение задаётся в шаблоне БП",
                )
        return render(
            request,
            "ui/specialist/template_preview.html",
            {
                "doc_type": doc_type,
                "preview_html": html,
                "requisites": requisites,
                "simple_preview": simple,
                "bpm_step_preview": bpm_step_preview,
            },
        )
